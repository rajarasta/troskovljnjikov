"""
Word (DOCX) Processing Skill - Create and analyze Word documents.

Provides tools for:
- Reading Word documents
- Creating Word documents
- Extracting text and formatting info
"""
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


def read_docx_text(file_path: str) -> dict:
    """
    Extract text from a Word document.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with extracted text and structure
    """
    if not HAS_PYTHON_DOCX:
        return {"error": "python-docx not installed. Install with: pip install python-docx"}

    try:
        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc = Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name,
                "level": para.paragraph_format.outline_level,
            })

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return {
            "file": str(path),
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "content": paragraphs,
            "tables_data": tables,
            "success": True,
        }
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return {"error": str(e), "success": False}


def create_docx_file(
    file_path: str,
    content: dict,
) -> dict:
    """
    Create a new Word document with specified content.

    Args:
        file_path: Path where to save the DOCX file
        content: Dictionary with document content:
            - title: Document title
            - paragraphs: List of paragraph strings or dicts
            - tables: List of table data (rows of cells)

    Returns:
        Dictionary with creation status
    """
    if not HAS_PYTHON_DOCX:
        return {"error": "python-docx not installed"}

    try:
        doc = Document()

        # Add title if provided
        if "title" in content:
            title = doc.add_heading(content["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add paragraphs
        if "paragraphs" in content:
            for para in content["paragraphs"]:
                if isinstance(para, dict):
                    p = doc.add_paragraph(para.get("text", ""))
                    if "style" in para:
                        p.style = para["style"]
                else:
                    doc.add_paragraph(para)

        # Add tables
        if "tables" in content:
            for table_data in content["tables"]:
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0]) if table_data else 0
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = "Light Grid Accent 1"

                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = str(cell_text)

        doc.save(file_path)
        return {
            "file": str(file_path),
            "success": True,
        }
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        return {"error": str(e), "success": False}


def extract_docx_metadata(file_path: str) -> dict:
    """
    Extract metadata from a Word document.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with document metadata
    """
    if not HAS_PYTHON_DOCX:
        return {"error": "python-docx not installed"}

    try:
        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        doc = Document(path)
        core_props = doc.core_properties

        return {
            "file": str(path),
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "comments": core_props.comments,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "success": True,
        }
    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {e}")
        return {"error": str(e), "success": False}


# Skill registration info
SKILL_INFO = {
    "name": "docx_handler",
    "description": "Create and analyze Word documents",
    "functions": [
        {
            "name": "read_docx_text",
            "description": "Extract text and structure from a Word document",
            "parameters": {
                "file_path": "Path to the DOCX file",
            },
        },
        {
            "name": "create_docx_file",
            "description": "Create a new Word document",
            "parameters": {
                "file_path": "Path where to save the file",
                "content": "Dictionary with title, paragraphs, and tables",
            },
        },
        {
            "name": "extract_docx_metadata",
            "description": "Get metadata from a Word document",
            "parameters": {
                "file_path": "Path to the DOCX file",
            },
        },
    ],
}
